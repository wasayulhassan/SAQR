"""
file_context.py
Extracts a compact, model-friendly summary of a file attached directly in
the Chat panel, so the chatbot can have a real conversation about it —
answer questions, spot issues, summarize sections, etc.

This is intentionally separate from analyzer.py's full_analysis(): that
pipeline is for the dedicated Analyze tab (stats/trends/anomalies/charts).
This one is lighter-weight and works across more file types (spreadsheets,
plain text, PDF, Word docs), trading completeness for something that fits
in a small hosted model's context window.
"""

import os

from . import analyzer

MAX_CONTENT_CHARS = 3000  # keep prompts small — some fallback models have ~4k token windows total

SUPPORTED_EXTENSIONS = {"csv", "xlsx", "xls", "txt", "pdf", "docx"}


def _truncate(text: str, max_chars: int = MAX_CONTENT_CHARS) -> tuple:
    text = text.strip()
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars], True


def _extract_spreadsheet(filepath: str, filename: str) -> dict:
    df = analyzer.load_file(filepath)
    stats = analyzer.summary_stats(df)
    sample = df.head(25).to_csv(index=False)
    content, truncated = _truncate(sample)

    meta = f"{stats['rows']} rows × {len(stats['columns'])} columns"
    return {
        "type": "spreadsheet",
        "filename": filename,
        "meta": meta,
        "content_text": (
            f"Columns: {', '.join(stats['columns'])}\n"
            f"Numeric columns: {', '.join(stats['numeric_columns']) or 'none'}\n"
            f"Missing values per column: {stats['missing_values']}\n\n"
            f"Sample rows (first 25, CSV format):\n{content}"
        ),
        "truncated": truncated,
    }


def _extract_txt(filepath: str, filename: str) -> dict:
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        raw = f.read()
    content, truncated = _truncate(raw)
    word_count = len(raw.split())
    return {
        "type": "text",
        "filename": filename,
        "meta": f"~{word_count} words",
        "content_text": content,
        "truncated": truncated,
    }


def _extract_pdf(filepath: str, filename: str) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(filepath)
    pages_text = []
    for page in reader.pages:
        pages_text.append(page.extract_text() or "")
    full_text = "\n".join(pages_text)
    content, truncated = _truncate(full_text)
    return {
        "type": "pdf",
        "filename": filename,
        "meta": f"{len(reader.pages)} page(s)",
        "content_text": content,
        "truncated": truncated,
    }


def _extract_docx(filepath: str, filename: str) -> dict:
    from docx import Document

    doc = Document(filepath)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    content, truncated = _truncate(full_text)
    return {
        "type": "docx",
        "filename": filename,
        "meta": f"{len(doc.paragraphs)} paragraph(s)",
        "content_text": content,
        "truncated": truncated,
    }


def extract_context(filepath: str, filename: str) -> dict:
    """Returns a dict: {type, filename, meta, content_text, truncated} or
    raises ValueError for unsupported types."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: .{ext}. Supported: "
            + ", ".join(sorted(SUPPORTED_EXTENSIONS))
        )

    if ext in ("csv", "xlsx", "xls"):
        return _extract_spreadsheet(filepath, filename)
    if ext == "txt":
        return _extract_txt(filepath, filename)
    if ext == "pdf":
        return _extract_pdf(filepath, filename)
    if ext == "docx":
        return _extract_docx(filepath, filename)

    raise ValueError(f"Unsupported file type: .{ext}")
