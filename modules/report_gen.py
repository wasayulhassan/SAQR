"""
report_gen.py
Builds a Word (.docx) report from an analysis dict produced by analyzer.py
"""

import os
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def build_report(analysis: dict, title: str = "Data Analysis Report") -> str:
    doc = Document()

    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = analysis["summary"]
    trends = analysis["trends"]
    anomalies = analysis["anomalies"]
    charts = analysis["charts"]

    # Overview
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(f"Rows: {summary['rows']}")
    doc.add_paragraph(f"Columns: {', '.join(summary['columns'])}")
    missing = {k: v for k, v in summary["missing_values"].items() if v > 0}
    if missing:
        doc.add_paragraph("Missing values detected: " +
                           ", ".join(f"{k} ({v})" for k, v in missing.items()))
    else:
        doc.add_paragraph("No missing values detected.")

    # Summary statistics table
    doc.add_heading("2. Summary Statistics", level=1)
    describe = summary["describe"]
    if describe:
        cols = list(describe.keys())
        stats_names = list(next(iter(describe.values())).keys())
        table = doc.add_table(rows=1, cols=len(cols) + 1)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Statistic"
        for i, c in enumerate(cols):
            hdr[i + 1].text = c
        for stat in stats_names:
            row = table.add_row().cells
            row[0].text = stat
            for i, c in enumerate(cols):
                row[i + 1].text = str(describe[c][stat])

    # Trends
    doc.add_heading("3. Trends", level=1)
    for col, t in trends.items():
        if isinstance(t, dict):
            doc.add_paragraph(
                f"• {col}: {t['direction']} trend "
                f"({t['pct_change_start_to_end']}% change start-to-end)",
                style="List Bullet",
            )
        else:
            doc.add_paragraph(f"• {col}: {t}", style="List Bullet")

    # Anomalies
    doc.add_heading("4. Anomalies", level=1)
    if anomalies:
        for col, info in anomalies.items():
            doc.add_paragraph(
                f"• {col}: {len(info['outlier_row_indices'])} outlier(s) "
                f"at rows {info['outlier_row_indices']}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("No significant anomalies detected.")

    # Charts
    if charts:
        doc.add_heading("5. Charts", level=1)
        for chart_path in charts:
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(5.5))

    fname = f"report_{uuid.uuid4().hex[:8]}.docx"
    fpath = os.path.join(OUTPUT_DIR, fname)
    doc.save(fpath)
    return fpath
